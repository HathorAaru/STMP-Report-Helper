from copy import deepcopy

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


HIGHLIGHT_FILL = "FFFF00"


def normalise_text(text):
    return " ".join(str(text or "").split()).casefold()


def table_text(table):
    return normalise_text(" ".join(cell.text for row in table.rows for cell in row.cells))


def find_table_containing(document, text):
    needle = normalise_text(text)
    for table in document.tables:
        if needle in table_text(table):
            return table
    raise ValueError(f"Could not find table containing: {text}")


def set_cell_text_preserve_first_run(cell, text):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = str(text)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(str(text))
    for paragraph in cell.paragraphs[1:]:
        for run in paragraph.runs:
            run.text = ""


def shade_cell(cell, fill=HIGHLIGHT_FILL):
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clear_cell_shading(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)


def clone_document_element(element):
    return deepcopy(element)

